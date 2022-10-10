from docx import Document
from docx.shared import Cm
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


document = Document()

document.add_picture('me.jpg', width=Cm(5.5), height=Cm(5))

name = input('Enter your name : ')
speak('Hello' + name + 'How you doing?')
phone = input('Enter your phone number : ')
email = input('Enter your email : ')

document.add_paragraph(
    name + ' | ' + phone + ' | ' + email
)

# about
document.add_heading('About me')
document.add_paragraph(input('Tell us about yourself : '))

# previous experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('The company name you previously worked for : ')
start_date = input('From date : ')
end_date = input('To date : ')
describe_experience = input('Describe your experience at ' + company + ' : ')
p.add_run(company + ' ').bold = True
p.add_run(start_date + ' - ' + end_date + '\n').italic = True
p.add_run(describe_experience)

while True:
    has_more_experience = input('Do you have more work experiences ? Yes / No ')
    if has_more_experience.lower() == 'yes':
        p = document.add_paragraph()
        company = input('The company name you previously worked for : ')
        start_date = input('From date : ')
        end_date = input('To date : ')
        describe_experience = input('Describe your experience at ' + company + ' : ')
        p.add_run(company + ' ').bold = True
        p.add_run(start_date + ' - ' + end_date + '\n').italic = True
        p.add_run(describe_experience)
    else:
        break
# skills
document.add_heading('Skills')
skills = input('Add your skill : ')
document.add_paragraph(skills, style='List Bullet')

while True:
    has_more_skills = input('Do you have more skills to add ? Yes/ No ')
    if has_more_skills.lower() == 'yes':
        skills = input('Add your skill : ')
        document.add_paragraph(skills, style='List Bullet')
    else:
        break

section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'Cv is programmed by Hulya'

document.save('cv.docx')
